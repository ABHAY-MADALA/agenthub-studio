"""
parsers.py
==========
File -> List[ParsedSection]. One function per supported file type, plus
parse_document() which picks the right one from the file extension.

PDFs keep their page boundaries (page=1, 2, 3...) so citations can say
"page 4". Everything else comes back as a single section (page=None) -
chunking.py is what actually splits it into retrievable pieces.

No OCR: if a PDF page has no extractable text (e.g. a scanned image),
that page is skipped and logged, not silently invented.
"""

from pathlib import Path
from typing import List

import config
from .models import EmptyDocumentError, ParsedSection, UnsupportedFileTypeError

logger = config.get_logger(__name__)


def parse_document(path: Path | str, file_type: str) -> List[ParsedSection]:
    path = Path(path)
    parser = _PARSERS.get(file_type)
    if parser is None:
        raise UnsupportedFileTypeError(f"'{file_type}' is not a supported file type.")

    sections = parser(path)
    sections = [s for s in sections if s.text and s.text.strip()]
    if not sections:
        raise EmptyDocumentError(
            "No extractable text was found in this file. Scanned/image-only PDFs aren't "
            "supported (no OCR)."
        )
    return sections


def _parse_txt(path: Path) -> List[ParsedSection]:
    text = _read_text_file(path)
    return [ParsedSection(text=text)]


def _parse_markdown(path: Path) -> List[ParsedSection]:
    # Markdown is kept as-is (not stripped to plain text) - the headings
    # and structure are useful context for both chunking and the LLM.
    text = _read_text_file(path)
    return [ParsedSection(text=text)]


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise EmptyDocumentError(f"Could not decode '{path.name}' as text.")


def _parse_pdf(path: Path) -> List[ParsedSection]:
    import pymupdf  # local import - keep the dependency optional at module load time

    sections: List[ParsedSection] = []
    with pymupdf.open(path) as doc:
        for page_index, page in enumerate(doc):
            text = page.get_text("text")
            if text and text.strip():
                sections.append(ParsedSection(text=text, page=page_index + 1))
            else:
                logger.info("PDF '%s' page %d has no extractable text (skipped, no OCR).", path.name, page_index + 1)
    return sections


def _parse_docx(path: Path) -> List[ParsedSection]:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts: List[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return [ParsedSection(text="\n\n".join(parts))]


def _parse_csv(path: Path) -> List[ParsedSection]:
    import pandas as pd

    df = pd.read_csv(path, dtype=str, keep_default_na=False, on_bad_lines="skip")
    if df.empty:
        return []

    columns = list(df.columns)
    sections: List[ParsedSection] = []

    # One section per batch of rows (not one giant blob) so chunking.py's
    # splitter has natural break points and metadata can report which
    # rows a chunk came from.
    batch_size = 200
    for start in range(0, len(df), batch_size):
        batch = df.iloc[start : start + batch_size]
        lines = [", ".join(f"{col}: {row[col]}" for col in columns) for _, row in batch.iterrows()]
        text = "\n".join(lines)
        sections.append(ParsedSection(
            text=text,
            extra={"row_start": int(start), "row_end": int(start + len(batch) - 1)},
        ))
    return sections


_PARSERS = {
    "txt": _parse_txt,
    "md": _parse_markdown,
    "markdown": _parse_markdown,
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "csv": _parse_csv,
}
